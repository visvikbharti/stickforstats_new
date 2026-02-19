"""
Manuscript Parser for Statistical Quality Review

This module parses academic manuscripts (PDF, LaTeX, DOCX) and extracts
structured sections for statistical quality analysis. It is designed to
integrate with the SQS (Statistical Quality Score) scoring pipeline as
part of StickForStats' journal integration platform.

Supported formats:
- PDF (via pdfplumber with PyPDF2 fallback)
- LaTeX (.tex)
- DOCX (via python-docx)

The parser segments manuscripts into academic sections (Abstract, Methods,
Results, etc.), extracts metadata, detects statistical test mentions, and
identifies figure/table references --- producing a ParsedManuscript that
downstream components (SQSScorer, Guardian) can consume directly.
"""

from __future__ import annotations

import io
import os
import re
import logging
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Tuple

# ---------------------------------------------------------------------------
# Graceful optional-dependency imports
# ---------------------------------------------------------------------------

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx as python_docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

# ===========================================================================
# Section detection patterns (case-insensitive)
# ===========================================================================

SECTION_PATTERNS: List[Tuple[str, str]] = [
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:abstract)\s*\n',
        'abstract',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:introduction|background)\s*\n',
        'introduction',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:methods?|materials?\s+and\s+methods?'
        r'|methodology|experimental\s+(?:design|methods?|procedures?))\s*\n',
        'methods',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:results?|findings?)\s*\n',
        'results',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:discussion|interpretation)\s*\n',
        'discussion',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:conclusions?|summary)\s*\n',
        'conclusion',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:references?|bibliography)\s*\n',
        'references',
    ),
    (
        r'(?:^|\n)\s*(?:\d+\.?\s+)?(?:supplementary|appendix|appendices)\s*\n',
        'supplementary',
    ),
]

# ===========================================================================
# Statistical test detection patterns
# ===========================================================================

TEST_PATTERNS: Dict[str, str] = {
    't-test': r"t[\s\-]?test|student'?s?\s+t",
    'ANOVA': r'ANOVA|analysis\s+of\s+variance',
    'chi-square': r'chi[\s\-]?square|\u03c7\u00b2|\u03c72',
    'regression': r'(?:linear|logistic|multiple)\s+regression',
    'correlation': r'(?:Pearson|Spearman|Kendall)\s+correlation',
    'Mann-Whitney': r'Mann[\s\-]?Whitney',
    'Kruskal-Wallis': r'Kruskal[\s\-]?Wallis',
    'Wilcoxon': r'Wilcoxon',
    'Fisher': r"Fisher'?s?\s+exact",
    'McNemar': r'McNemar',
    'Friedman': r'Friedman',
    'mixed-effects': r'mixed[\s\-]?(?:effects?|model)',
    'Bayesian': r'Bayesian\s+(?:analysis|inference|statistics)',
    'survival': (
        r'(?:Cox|Kaplan[\s\-]?Meier|survival)'
        r'\s+(?:analysis|regression|model)'
    ),
    'SEM': r'structural\s+equation\s+model',
    'factor-analysis': r'(?:exploratory|confirmatory)\s+factor\s+analysis',
}

# ===========================================================================
# Journal-format detection patterns
# ===========================================================================

_JOURNAL_FORMAT_PATTERNS: Dict[str, str] = {
    'APA': (
        r'APA\s+(?:style|format|guidelines|7th|6th)'
        r'|[Ff]\s*\(\s*\d+\s*,\s*\d+\s*\)\s*=\s*[\d.]+\s*,\s*p\s*[=<]'
    ),
    'CONSORT': r'CONSORT|flow\s+diagram',
    'STROBE': r'STROBE',
    'PRISMA': r'PRISMA',
    'JARS-Quant': r'JARS|Journal\s+Article\s+Reporting\s+Standards',
    'ARRIVE': r'ARRIVE\s+guidelines?',
    'EQUATOR': r'EQUATOR',
}

# ===========================================================================
# LaTeX command-stripping helpers
# ===========================================================================

# Commands whose single brace-delimited argument should be kept as plain text
_LATEX_UNWRAP_COMMANDS = re.compile(
    r'\\(?:textbf|textit|texttt|emph|underline|textsc|textrm|textsf)\{([^}]*)\}'
)

# Commands that should be replaced with a readable token
_LATEX_CITE = re.compile(r'\\(?:cite[tp]?|citep|citet|citeauthor)\{[^}]*\}')

# Whole-line comment
_LATEX_COMMENT = re.compile(r'(?m)^\s*%.*$')

# Inline comment (not preceded by backslash)
_LATEX_INLINE_COMMENT = re.compile(r'(?<!\\)%.*$', re.MULTILINE)

# Environment markers — keep math, strip the rest
_LATEX_ENV_BEGIN = re.compile(r'\\begin\{(?!(?:equation|align|math|eqnarray|gather|multline)\*?\})[^}]*\}')
_LATEX_ENV_END = re.compile(r'\\end\{(?!(?:equation|align|math|eqnarray|gather|multline)\*?\})[^}]*\}')

# Generic remaining commands (e.g. \label{...}, \ref{...}, \usepackage{...})
_LATEX_GENERIC_CMD = re.compile(r'\\(?:label|ref|usepackage|bibliographystyle|bibliography|documentclass'
                                r'|author|date|maketitle|tableofcontents|newcommand|renewcommand'
                                r'|include|input|vspace|hspace|noindent|centering|pagestyle'
                                r'|setlength|addtolength)\b(?:\[[^\]]*\])?(?:\{[^}]*\})*')

# Section-like commands for heading extraction
_LATEX_SECTION_CMD = re.compile(
    r'\\(?:section|subsection|subsubsection|chapter|paragraph|subparagraph)\*?\{([^}]*)\}'
)

# \title{...} extraction
_LATEX_TITLE_CMD = re.compile(r'\\title\{([^}]*)\}')

# \author{...} extraction
_LATEX_AUTHOR_CMD = re.compile(r'\\author\{([^}]*)\}')

# \keywords{...} extraction
_LATEX_KEYWORDS_CMD = re.compile(r'\\keywords\{([^}]*)\}')


# ===========================================================================
# Data classes
# ===========================================================================

@dataclass
class ManuscriptSection:
    """A single logical section extracted from a manuscript."""

    section_type: str       # abstract, introduction, methods, results, ...
    title: str              # detected heading text
    content: str            # full text content of the section
    start_pos: int          # character position in full text
    end_pos: int            # character position in full text
    tables: List[Dict[str, Any]] = field(default_factory=list)
    figures_mentioned: List[str] = field(default_factory=list)


@dataclass
class ManuscriptMetadata:
    """Top-level metadata extracted from a manuscript."""

    title: Optional[str]
    authors: List[str]
    abstract: Optional[str]
    keywords: List[str]
    word_count: int
    page_count: int
    sections_found: List[str]
    statistical_tests_mentioned: List[str]
    journal_format: Optional[str]
    raw_text: str


@dataclass
class ParsedManuscript:
    """
    Complete result of parsing a manuscript.

    Provides structured sections, metadata, and convenience accessors for
    the methods and results text that the SQS pipeline consumes most often.
    """

    metadata: ManuscriptMetadata
    sections: List[ManuscriptSection]
    full_text: str
    methods_text: str       # convenience: extracted methods section (empty if absent)
    results_text: str       # convenience: extracted results section (empty if absent)
    parse_quality: float    # 0-1, how well we could parse the document
    warnings: List[str]     # parsing issues encountered


# ===========================================================================
# Main parser
# ===========================================================================

class ManuscriptParser:
    """
    Parse academic manuscripts (PDF, LaTeX, DOCX) into structured sections
    suitable for downstream Statistical Quality Score analysis.

    Usage::

        parser = ManuscriptParser()
        result = parser.parse(uploaded_file, file_type='auto')

    For LaTeX files, metadata (title, authors, keywords) is extracted from
    the raw source *before* command stripping so that ``\\title{...}`` etc.
    are not lost.

    ::

        scorer = SQSScorer()
        report = scorer.analyze(result.full_text)
    """

    # Supported file extensions mapped to canonical type key
    _EXT_MAP: Dict[str, str] = {
        '.pdf': 'pdf',
        '.tex': 'latex',
        '.latex': 'latex',
        '.docx': 'docx',
    }

    def __init__(self) -> None:
        # Holds metadata extracted from raw LaTeX source before stripping.
        # Populated by _extract_text_latex(); consumed by _extract_metadata().
        self._latex_metadata: Dict[str, Any] = {}

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def parse(self, file: Any, file_type: str = 'auto') -> ParsedManuscript:
        """
        Parse a manuscript file and return structured content.

        Args:
            file: A file-like object (e.g. Django ``UploadedFile``,
                  ``io.BytesIO``, or an open file handle).  For LaTeX
                  files the content may also be passed as ``io.StringIO``.
            file_type: One of ``'pdf'``, ``'latex'``, ``'docx'``, or
                       ``'auto'`` (default) to detect from the file name
                       or magic bytes.

        Returns:
            A :class:`ParsedManuscript` instance with sections, metadata,
            and quality indicators.

        Raises:
            ValueError: If the file type cannot be determined or no
                extraction library is available.
        """
        warnings: List[str] = []

        # 1. Determine file type
        if file_type == 'auto':
            file_type = self._detect_file_type(file)
            if file_type == 'unknown':
                raise ValueError(
                    'Could not determine file type. '
                    'Please specify file_type explicitly.'
                )

        # 2. Extract raw text
        text, page_count, extract_warnings = self._extract_text(
            file, file_type
        )
        warnings.extend(extract_warnings)

        if not text or len(text.strip()) < 50:
            warnings.append(
                'Extracted text is very short; parsing quality will be low.'
            )

        # 3. Segment into sections
        sections = self._segment_sections(text)

        # 4. Enrich sections with tables and figure references
        for section in sections:
            section.tables = self._extract_tables_from_section(section.content)
            section.figures_mentioned = self._detect_figure_references(
                section.content
            )

        # 5. Extract metadata
        metadata = self._extract_metadata(text, sections, page_count)

        # 6. Convenience accessors
        methods_text = ''
        results_text = ''
        for sec in sections:
            if sec.section_type == 'methods' and not methods_text:
                methods_text = sec.content
            elif sec.section_type == 'results' and not results_text:
                results_text = sec.content

        # 7. Compute parse quality (0-1)
        parse_quality = self._compute_parse_quality(text, sections, warnings)

        return ParsedManuscript(
            metadata=metadata,
            sections=sections,
            full_text=text,
            methods_text=methods_text,
            results_text=results_text,
            parse_quality=parse_quality,
            warnings=warnings,
        )

    # ------------------------------------------------------------------
    # File-type detection
    # ------------------------------------------------------------------

    def _detect_file_type(self, file: Any) -> str:
        """
        Auto-detect file type from extension or magic bytes.

        Args:
            file: A file-like object (may have a ``name`` attribute).

        Returns:
            One of ``'pdf'``, ``'latex'``, ``'docx'``, or ``'unknown'``.
        """
        # Try extension first
        name = getattr(file, 'name', '') or ''
        if name:
            ext = os.path.splitext(name)[1].lower()
            if ext in self._EXT_MAP:
                return self._EXT_MAP[ext]

        # Fall back to magic bytes
        try:
            pos = file.tell()
            header = file.read(8)
            file.seek(pos)

            if isinstance(header, str):
                header = header.encode('utf-8', errors='ignore')

            # PDF magic: %PDF
            if header[:4] == b'%PDF':
                return 'pdf'

            # DOCX (ZIP with PK signature)
            if header[:2] == b'PK':
                return 'docx'

            # LaTeX heuristic: starts with \ or %
            if header[:1] in (b'\\', b'%'):
                return 'latex'

        except Exception:
            pass

        return 'unknown'

    # ------------------------------------------------------------------
    # Text extraction dispatchers
    # ------------------------------------------------------------------

    def _extract_text(
        self, file: Any, file_type: str
    ) -> Tuple[str, int, List[str]]:
        """
        Dispatch to the correct extractor and return ``(text, page_count, warnings)``.
        """
        if file_type == 'pdf':
            return self._extract_text_pdf(file)
        elif file_type == 'latex':
            return self._extract_text_latex(file)
        elif file_type == 'docx':
            return self._extract_text_docx(file)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    # ------------------------------------------------------------------
    # PDF extraction (pdfplumber primary, PyPDF2 fallback)
    # ------------------------------------------------------------------

    def _extract_text_pdf(self, file: Any) -> Tuple[str, int, List[str]]:
        """
        Extract text from a PDF file.

        Uses pdfplumber for higher-quality extraction with a PyPDF2
        fallback, mirroring the pattern in ``sqs_views.extract_text_from_pdf``.

        Args:
            file: File-like object containing PDF data.

        Returns:
            Tuple of ``(text, page_count, warnings)``.
        """
        text = ''
        page_count = 0
        warnings: List[str] = []

        # --- pdfplumber (preferred) ---
        if PDFPLUMBER_AVAILABLE:
            try:
                file.seek(0)
                with pdfplumber.open(file) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                if text.strip():
                    return text, page_count, warnings
            except Exception as exc:
                logger.warning('pdfplumber extraction failed: %s', exc)
                warnings.append(f'pdfplumber extraction failed: {exc}')

        # --- PyPDF2 fallback ---
        if PYPDF2_AVAILABLE:
            try:
                file.seek(0)
                reader = PyPDF2.PdfReader(file)
                page_count = len(reader.pages)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + '\n'
                if text.strip():
                    return text, page_count, warnings
            except Exception as exc:
                logger.warning('PyPDF2 extraction failed: %s', exc)
                warnings.append(f'PyPDF2 extraction failed: {exc}')

        if not text.strip():
            warnings.append(
                'No PDF extraction library available or extraction produced '
                'no text. Install pdfplumber or PyPDF2.'
            )

        return text, page_count, warnings

    # ------------------------------------------------------------------
    # LaTeX extraction
    # ------------------------------------------------------------------

    def _extract_text_latex(self, file: Any) -> Tuple[str, int, List[str]]:
        """
        Extract readable text from a LaTeX source file.

        Strips LaTeX commands while preserving ``$...$`` math notation
        and converting ``\\section{Title}`` into plain-text headings.

        Args:
            file: File-like object containing ``.tex`` data.

        Returns:
            Tuple of ``(text, page_count, warnings)``.
            ``page_count`` is estimated as ``ceil(word_count / 250)``.
        """
        warnings: List[str] = []

        try:
            file.seek(0)
            raw = file.read()
            if isinstance(raw, bytes):
                raw = raw.decode('utf-8', errors='replace')
        except Exception as exc:
            logger.warning('LaTeX read failed: %s', exc)
            warnings.append(f'Could not read LaTeX file: {exc}')
            return '', 0, warnings

        # Extract metadata from raw source *before* stripping commands,
        # because \title{}, \author{}, \keywords{} will be removed.
        self._latex_metadata = self._extract_latex_metadata_from_source(raw)

        text = self._strip_latex(raw)

        # Rough page estimate
        word_count = len(text.split())
        page_count = max(1, -(-word_count // 250))  # ceiling division

        return text, page_count, warnings

    @staticmethod
    def _extract_latex_metadata_from_source(raw: str) -> Dict[str, Any]:
        """
        Pull title, authors, and keywords from raw LaTeX source before
        the stripping pass destroys the commands.
        """
        meta: Dict[str, Any] = {}

        m = _LATEX_TITLE_CMD.search(raw)
        if m:
            meta['title'] = m.group(1).strip()

        m = _LATEX_AUTHOR_CMD.search(raw)
        if m:
            parts = re.split(r'\\and|,\s*|\s+and\s+', m.group(1))
            meta['authors'] = [p.strip() for p in parts if p.strip()]

        m = _LATEX_KEYWORDS_CMD.search(raw)
        if m:
            meta['keywords'] = [
                kw.strip() for kw in re.split(r'[;,]', m.group(1)) if kw.strip()
            ]

        return meta

    def _strip_latex(self, raw: str) -> str:
        """
        Convert LaTeX source to plain text.

        Strategy:
        1. Remove comments (``%`` lines).
        2. Convert ``\\section{...}`` to ``\\nTitle\\n`` headings.
        3. Unwrap formatting commands (``\\textbf``, ``\\emph``, etc.).
        4. Replace ``\\cite{...}`` with ``[citation]``.
        5. Remove non-math environments and generic commands.
        6. Preserve ``$...$`` and ``$$...$$`` math notation as-is.
        7. Collapse excess whitespace.
        """
        text = raw

        # Strip comments
        text = _LATEX_INLINE_COMMENT.sub('', text)

        # Convert section commands to headings
        text = _LATEX_SECTION_CMD.sub(r'\n\1\n', text)

        # Unwrap formatting commands
        text = _LATEX_UNWRAP_COMMANDS.sub(r'\1', text)

        # Replace citations
        text = _LATEX_CITE.sub('[citation]', text)

        # Remove non-math environments
        text = _LATEX_ENV_BEGIN.sub('', text)
        text = _LATEX_ENV_END.sub('', text)

        # Remove generic commands
        text = _LATEX_GENERIC_CMD.sub('', text)

        # Remove remaining simple commands (e.g. \\ , \newline, \item)
        text = re.sub(r'\\(?:newline|item|\\)\b', '\n', text)

        # Remove leftover backslash-commands that were not caught above,
        # but do NOT touch math ($...$).  We keep it simple: remove
        # \command{arg} and bare \command sequences outside math.
        text = re.sub(r'\\[a-zA-Z]+\*?(?:\[[^\]]*\])?(?:\{[^}]*\})*', '', text)

        # Collapse whitespace (preserve paragraph breaks)
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = text.strip()

        return text

    # ------------------------------------------------------------------
    # DOCX extraction
    # ------------------------------------------------------------------

    def _extract_text_docx(self, file: Any) -> Tuple[str, int, List[str]]:
        """
        Extract text from a DOCX file using python-docx.

        Args:
            file: File-like object containing DOCX data.

        Returns:
            Tuple of ``(text, page_count, warnings)``.
            ``page_count`` is estimated as ``ceil(word_count / 250)``.
        """
        warnings: List[str] = []

        if not DOCX_AVAILABLE:
            warnings.append(
                'python-docx is not installed. '
                'Install it with: pip install python-docx'
            )
            return '', 0, warnings

        try:
            file.seek(0)
            doc = python_docx.Document(file)

            paragraphs: List[str] = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            full_text = '\n'.join(paragraphs)

            # Rough page estimate
            word_count = len(full_text.split())
            page_count = max(1, -(-word_count // 250))

            return full_text, page_count, warnings

        except Exception as exc:
            logger.warning('DOCX extraction failed: %s', exc)
            warnings.append(f'DOCX extraction failed: {exc}')
            return '', 0, warnings

    # ------------------------------------------------------------------
    # Section segmentation
    # ------------------------------------------------------------------

    def _segment_sections(self, text: str) -> List[ManuscriptSection]:
        """
        Detect section breaks in the full text and segment accordingly.

        Uses regex patterns for common academic headings. When no known
        headings are found the entire text is returned as a single
        ``unknown`` section.

        Args:
            text: Full extracted manuscript text.

        Returns:
            List of :class:`ManuscriptSection` instances ordered by
            position in the text.
        """
        if not text:
            return []

        # Collect all heading matches: (start_pos, end_of_heading, type, title)
        matches: List[Tuple[int, int, str, str]] = []

        for pattern, section_type in SECTION_PATTERNS:
            for m in re.finditer(pattern, text, re.IGNORECASE):
                heading_text = m.group(0).strip()
                matches.append((m.start(), m.end(), section_type, heading_text))

        # Also detect LaTeX-style headings that survived stripping
        # (e.g. a raw "Introduction" on its own line)
        # Already handled by SECTION_PATTERNS above.

        if not matches:
            # No recognised sections -- return the whole document as one block
            return [
                ManuscriptSection(
                    section_type='unknown',
                    title='',
                    content=text,
                    start_pos=0,
                    end_pos=len(text),
                )
            ]

        # Sort by position
        matches.sort(key=lambda m: m[0])

        sections: List[ManuscriptSection] = []

        # If there is text before the first detected heading, treat it as
        # a preamble (possibly containing title/abstract).
        first_start = matches[0][0]
        if first_start > 0:
            preamble_text = text[:first_start].strip()
            if preamble_text:
                sections.append(
                    ManuscriptSection(
                        section_type='unknown',
                        title='',
                        content=preamble_text,
                        start_pos=0,
                        end_pos=first_start,
                    )
                )

        for idx, (start, heading_end, stype, title) in enumerate(matches):
            # Content runs from end-of-heading to start-of-next-heading
            if idx + 1 < len(matches):
                content_end = matches[idx + 1][0]
            else:
                content_end = len(text)

            content = text[heading_end:content_end].strip()

            sections.append(
                ManuscriptSection(
                    section_type=stype,
                    title=title,
                    content=content,
                    start_pos=start,
                    end_pos=content_end,
                )
            )

        return sections

    # ------------------------------------------------------------------
    # Metadata extraction
    # ------------------------------------------------------------------

    def _extract_metadata(
        self,
        text: str,
        sections: List[ManuscriptSection],
        page_count: int,
    ) -> ManuscriptMetadata:
        """
        Extract manuscript metadata from the full text and parsed sections.

        Args:
            text: Full extracted text.
            sections: Already-parsed sections.
            page_count: Page count from the extraction step.

        Returns:
            :class:`ManuscriptMetadata` instance.
        """
        title = self._extract_title(text)
        authors = self._extract_authors(text)
        abstract = self._extract_abstract(text, sections)
        keywords = self._extract_keywords(text)
        stat_tests = self._detect_statistical_tests(text)
        journal_format = self._detect_journal_format(text)

        sections_found = [s.section_type for s in sections]
        word_count = len(text.split())

        return ManuscriptMetadata(
            title=title,
            authors=authors,
            abstract=abstract,
            keywords=keywords,
            word_count=word_count,
            page_count=page_count,
            sections_found=sections_found,
            statistical_tests_mentioned=stat_tests,
            journal_format=journal_format,
            raw_text=text,
        )

    def _extract_title(self, text: str) -> Optional[str]:
        """
        Extract the manuscript title.

        Consults pre-extracted LaTeX metadata first, then tries
        ``\\title{...}`` in the text, then falls back to the first
        non-empty line.
        """
        # Pre-extracted from raw LaTeX source (before stripping)
        if self._latex_metadata.get('title'):
            return self._latex_metadata['title']

        # LaTeX title in text (may still work for non-stripped input)
        m = _LATEX_TITLE_CMD.search(text)
        if m:
            return m.group(1).strip()

        # Heuristic: first non-empty line (often the title in extracted PDFs)
        for line in text.split('\n'):
            line = line.strip()
            if line and len(line) > 5:
                # Avoid picking up boilerplate / page numbers
                if not re.match(r'^\d+$', line):
                    return line
        return None

    def _extract_authors(self, text: str) -> List[str]:
        """
        Extract author names.

        Consults pre-extracted LaTeX metadata first, then tries
        ``\\author{...}`` in the text, then falls back to heuristic
        name detection near the top of the document.
        """
        # Pre-extracted from raw LaTeX source
        if self._latex_metadata.get('authors'):
            return self._latex_metadata['authors']

        authors: List[str] = []

        # LaTeX author command in text
        m = _LATEX_AUTHOR_CMD.search(text)
        if m:
            raw = m.group(1)
            # Split on \\and, commas, or "and"
            parts = re.split(r'\\and|,\s*|\s+and\s+', raw)
            for part in parts:
                name = part.strip()
                if name and len(name) > 1:
                    authors.append(name)
            if authors:
                return authors

        # Heuristic: look in the first ~500 chars for name-like patterns
        # (two or more capitalised words separated by commas or "and")
        header = text[:500]
        name_pattern = re.compile(
            r'([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)'
        )
        found = name_pattern.findall(header)
        if found:
            # Deduplicate while preserving order
            seen: set[str] = set()
            for name in found:
                name = name.strip()
                if name not in seen:
                    seen.add(name)
                    authors.append(name)

        return authors

    def _extract_abstract(
        self, text: str, sections: List[ManuscriptSection]
    ) -> Optional[str]:
        """Return the abstract text if an abstract section was found."""
        for sec in sections:
            if sec.section_type == 'abstract':
                return sec.content

        # Fallback: look for "Abstract" keyword followed by text
        m = re.search(
            r'(?:^|\n)\s*abstract[:\s]*\n?(.*?)(?=\n\s*(?:introduction|background|keywords?|1[\.\s]))',
            text,
            re.IGNORECASE | re.DOTALL,
        )
        if m:
            return m.group(1).strip()

        return None

    def _extract_keywords(self, text: str) -> List[str]:
        """
        Extract keywords from the manuscript.

        Consults pre-extracted LaTeX metadata first, then tries
        ``\\keywords{...}`` in the text, then plain-text patterns.
        """
        # Pre-extracted from raw LaTeX source
        if self._latex_metadata.get('keywords'):
            return self._latex_metadata['keywords']

        # LaTeX command in text
        m = _LATEX_KEYWORDS_CMD.search(text)
        if m:
            raw = m.group(1)
        else:
            # Plain text: "Keywords: word1, word2, ..."
            m = re.search(
                r'(?:keywords?|key\s+words?)\s*[:\-]\s*(.+?)(?:\n\n|\n[A-Z1-9])',
                text,
                re.IGNORECASE,
            )
            if m:
                raw = m.group(1)
            else:
                return []

        # Split on commas or semicolons
        keywords = [kw.strip() for kw in re.split(r'[;,]', raw) if kw.strip()]
        return keywords

    # ------------------------------------------------------------------
    # Statistical test detection
    # ------------------------------------------------------------------

    def _detect_statistical_tests(self, text: str) -> List[str]:
        """
        Detect which statistical tests are mentioned in the text.

        Args:
            text: Full manuscript text.

        Returns:
            Sorted list of unique test name strings that were found.
        """
        found: List[str] = []
        for test_name, pattern in TEST_PATTERNS.items():
            if re.search(pattern, text, re.IGNORECASE):
                found.append(test_name)
        found.sort()
        return found

    # ------------------------------------------------------------------
    # Journal-format detection
    # ------------------------------------------------------------------

    def _detect_journal_format(self, text: str) -> Optional[str]:
        """
        Attempt to detect the reporting guideline / journal format.

        Returns the first match among known guideline patterns, or
        ``None`` if no guideline is detected.
        """
        for fmt, pattern in _JOURNAL_FORMAT_PATTERNS.items():
            if re.search(pattern, text, re.IGNORECASE):
                return fmt
        return None

    # ------------------------------------------------------------------
    # Table extraction from section text
    # ------------------------------------------------------------------

    def _extract_tables_from_section(
        self, section_text: str
    ) -> List[Dict[str, Any]]:
        """
        Detect table-like structures in a section's text.

        This is a best-effort heuristic that looks for:
        - Lines with multiple tab/multi-space separated columns
        - Patterns like "Table N." captions

        Args:
            section_text: Text content of a single section.

        Returns:
            List of dicts, each with ``'caption'`` and ``'rows'`` keys.
        """
        tables: List[Dict[str, Any]] = []

        # Detect "Table N" captions
        caption_pattern = re.compile(
            r'(Table\s+\d+[.:]\s*.+?)(?:\n|$)', re.IGNORECASE
        )

        for cap_match in caption_pattern.finditer(section_text):
            caption = cap_match.group(1).strip()

            # Try to collect rows after the caption that look tabular
            # (two or more "columns" separated by 2+ spaces or tabs)
            start = cap_match.end()
            rows: List[List[str]] = []

            remaining = section_text[start:start + 2000]  # limit scan
            for line in remaining.split('\n'):
                line = line.strip()
                if not line:
                    if rows:
                        break  # blank line after rows → end of table
                    continue

                # Split on two-or-more spaces or tabs
                cols = re.split(r'\t|  +', line)
                if len(cols) >= 2:
                    rows.append(cols)
                elif rows:
                    break  # non-tabular line after rows → end of table

            tables.append({
                'caption': caption,
                'rows': rows,
            })

        return tables

    # ------------------------------------------------------------------
    # Figure / table reference detection
    # ------------------------------------------------------------------

    def _detect_figure_references(self, text: str) -> List[str]:
        """
        Find references to figures and tables (e.g. "Figure 1", "Fig. 2",
        "Table 3") within a block of text.

        Args:
            text: Text to scan.

        Returns:
            Sorted list of unique reference strings.
        """
        pattern = re.compile(
            r'(?:Figure|Fig\.?|Table|Appendix)\s+[A-Z]?\d+',
            re.IGNORECASE,
        )
        matches = pattern.findall(text)

        # Normalise and deduplicate
        seen: set[str] = set()
        refs: List[str] = []
        for ref in matches:
            normalised = ref.strip()
            if normalised not in seen:
                seen.add(normalised)
                refs.append(normalised)

        refs.sort(key=lambda r: (r.split()[0].lower(), r))
        return refs

    # ------------------------------------------------------------------
    # Parse-quality estimation
    # ------------------------------------------------------------------

    def _compute_parse_quality(
        self,
        text: str,
        sections: List[ManuscriptSection],
        warnings: List[str],
    ) -> float:
        """
        Estimate how well the document was parsed on a 0-1 scale.

        Factors:
        - Sufficient text length
        - Number of recognised sections
        - Presence of methods and results
        - Absence of extraction warnings
        """
        if not text:
            return 0.0

        score = 0.0
        max_score = 0.0

        # 1. Text length (up to 0.2)
        max_score += 0.2
        word_count = len(text.split())
        if word_count >= 3000:
            score += 0.2
        elif word_count >= 1000:
            score += 0.15
        elif word_count >= 300:
            score += 0.1
        else:
            score += 0.05

        # 2. Recognised sections (up to 0.3)
        max_score += 0.3
        known_types = {
            s.section_type for s in sections if s.section_type != 'unknown'
        }
        # Typical paper has 5-7 recognised sections
        section_ratio = min(len(known_types) / 5.0, 1.0)
        score += 0.3 * section_ratio

        # 3. Methods section present (0.15)
        max_score += 0.15
        if any(s.section_type == 'methods' for s in sections):
            score += 0.15

        # 4. Results section present (0.15)
        max_score += 0.15
        if any(s.section_type == 'results' for s in sections):
            score += 0.15

        # 5. No warnings penalty (up to 0.2)
        max_score += 0.2
        if not warnings:
            score += 0.2
        elif len(warnings) == 1:
            score += 0.1
        # else: 0

        return round(min(score / max_score, 1.0) if max_score > 0 else 0.0, 2)
