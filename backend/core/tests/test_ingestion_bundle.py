"""
Tests for multi-format, multi-file (bundle) ingestion.
======================================================

Covers the editor/publisher upload case added 2026-06-27:
  - DOCX table-cell extraction (statistics in Word tables were silently dropped)
  - PDF table-cell reconstruction
  - JATS/NLM XML on upload (references excluded)
  - figure-image OCR (TIFF/PNG/JPEG) — gated on a tesseract binary
  - the bundle ingestion service (classify + parse + OCR + load tables + combine)
  - POST /api/v1/verify/bundle/ end to end

See docs/INGESTION_ARCHITECTURE.md.
"""

import io
import unittest

import numpy as np
from scipy import stats

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import SimpleTestCase, override_settings
from rest_framework import status
from rest_framework.test import APITestCase

from core.manuscript.parser import ManuscriptParser
from core.manuscript import image_ocr
from core.manuscript.bundle_ingest import ingest_bundle, make_multitable_linker


# --------------------------------------------------------------------------- helpers
def _docx_bytes(paragraph: str, table_rows=None) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_paragraph(paragraph)
    if table_rows:
        t = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row in enumerate(table_rows):
            for c, val in enumerate(row):
                t.cell(r, c).text = val
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_with_table_bytes() -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    buf = io.BytesIO()
    docp = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    data = [["Comparison", "Statistic", "p"], ["A vs B", "t(28) = 2.50", "p = 0.002"]]
    tbl = Table(data)
    tbl.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black)]))
    docp.build([Paragraph("Results.", styles["Normal"]), Spacer(1, 12), tbl])
    return buf.getvalue()


def _png_with_text_bytes(text: str) -> bytes:
    from PIL import Image, ImageDraw, ImageFont

    img = Image.new("RGB", (760, 90), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial.ttf", 30)
    except Exception:
        font = ImageFont.load_default()
    draw.text((15, 30), text, fill="black", font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return buf.getvalue()


def _two_group_csv_and_t():
    rng = np.random.default_rng(7)
    treat = rng.normal(58.0, 8.0, 40)
    ctrl = rng.normal(50.0, 8.0, 40)
    rows = ["group,biomarker"] + [f"treatment,{v}" for v in treat] + [f"control,{v}" for v in ctrl]
    t, p = stats.ttest_ind(treat, ctrl, equal_var=True)
    return ("\n".join(rows) + "\n").encode(), abs(float(t)), float(p)


_JATS = b"""<?xml version="1.0"?>
<article xmlns:xlink="http://www.w3.org/1999/xlink">
<front><article-meta><title-group><article-title>A study</article-title></title-group>
<abstract><p>Tested.</p></abstract></article-meta></front>
<body><sec><title>Results</title><p>Significant effect, t(30) = 3.10, p = 0.004.</p></sec></body>
<back><ref-list><ref><mixed-citation>Ref t(99)=9.9, p=0.9</mixed-citation></ref></ref-list></back>
</article>"""


# --------------------------------------------------------------------------- parser
class ParserExtractionTest(SimpleTestCase):
    def test_docx_table_cells_extracted(self):
        b = _docx_bytes("Results.", table_rows=[["Contrast", "Result"], ["A vs B", "t(28) = 2.50, p = 0.002"]])
        f = io.BytesIO(b)
        f.name = "m.docx"
        parsed = ManuscriptParser().parse(f, file_type="docx")
        self.assertIn("t(28) = 2.50", parsed.full_text)
        self.assertIn("p = 0.002", parsed.full_text)

    def test_pdf_table_cells_present(self):
        f = io.BytesIO(_pdf_with_table_bytes())
        f.name = "m.pdf"
        parsed = ManuscriptParser().parse(f, file_type="pdf")
        self.assertIn("t(28) = 2.50", parsed.full_text)

    def test_jats_parsed_and_references_excluded(self):
        f = io.BytesIO(_JATS)
        f.name = "a.xml"
        parsed = ManuscriptParser().parse(f, file_type="auto")  # magic-byte detect
        self.assertIn("t(30) = 3.10", parsed.full_text)
        self.assertNotIn("t(99)=9.9", parsed.full_text)  # reference list dropped

    def test_no_table_duplication_when_linear_text_has_it(self):
        # a stat appearing cleanly in linear PDF text must not be doubled by table reconstruction
        f = io.BytesIO(_pdf_with_table_bytes())
        f.name = "m.pdf"
        parsed = ManuscriptParser().parse(f, file_type="pdf")
        self.assertEqual(parsed.full_text.count("t(28) = 2.50"), 1)


# --------------------------------------------------------------------------- OCR
class ImageOcrTest(SimpleTestCase):
    def test_capabilities_shape(self):
        caps = image_ocr.ocr_capabilities()
        for k in ("pytesseract", "pillow", "tesseract_binary", "pdf2image", "scanned_pdf_ocr"):
            self.assertIn(k, caps)

    def test_garbage_image_degrades_gracefully(self):
        text, warns = image_ocr.ocr_image(io.BytesIO(b"not an image"))
        self.assertEqual(text, "")
        self.assertTrue(warns)  # a warning, never an exception

    @unittest.skipUnless(image_ocr.tesseract_available(), "tesseract binary not installed")
    def test_ocr_recovers_statistic(self):
        text, _ = image_ocr.ocr_image(io.BytesIO(_png_with_text_bytes("r = 0.61, p = 0.01")))
        self.assertIn("0.61", text)
        self.assertIn("0.01", text)


# --------------------------------------------------------------------------- bundle service
class BundleIngestServiceTest(SimpleTestCase):
    def test_classify_combine_and_load(self):
        csv, _, _ = _two_group_csv_and_t()
        items = [
            {"name": "m.docx", "kind": "manuscript", "detail": "docx",
             "content": _docx_bytes("A t-test, t(28) = 2.50, p = 0.002.",
                                    table_rows=[["x", "y"], ["A vs B", "F(2,45) = 6.10, p = 0.004"]])},
            {"name": "data.csv", "kind": "data", "detail": ".csv", "content": csv},
            {"name": "junk.zip", "kind": "unknown", "detail": "", "content": b"PK\x03\x04junk"},
        ]
        b = ingest_bundle(items)
        self.assertIn("t(28) = 2.50", b.manuscript_text)
        self.assertIn("F(2,45) = 6.10", b.manuscript_text)  # DOCX table
        self.assertEqual(len(b.dataframes), 1)
        rep = b.report()
        self.assertEqual(rep["n_files"], 3)
        self.assertEqual(rep["n_data_files"], 1)
        self.assertEqual(rep["n_unknown"], 1)

    def test_multitable_linker_none_without_data(self):
        self.assertIsNone(make_multitable_linker([]))

    @unittest.skipUnless(image_ocr.tesseract_available(), "tesseract binary not installed")
    def test_image_text_included(self):
        items = [
            {"name": "p.txt", "kind": "manuscript", "detail": "latex", "content": b"Results follow."},
            {"name": "fig.png", "kind": "image", "detail": ".png",
             "content": _png_with_text_bytes("r = 0.61, p = 0.01")},
        ]
        b = ingest_bundle(items)
        self.assertTrue(b.ocr_used)
        self.assertIn("0.61", b.manuscript_text)


# --------------------------------------------------------------------------- API
@override_settings(SECURE_SSL_REDIRECT=False)
class VerifyBundleAPITest(APITestCase):
    URL = "/api/v1/verify/bundle/"

    def test_bundle_endpoint_verifies_across_files(self):
        csv, t_abs, p = _two_group_csv_and_t()
        paper = ("Results. Serum biomarker was higher in the treatment group than the control "
                 f"group (t(78) = {t_abs:.2f}, p = {p:.3f}).")
        docx = SimpleUploadedFile("m.docx", _docx_bytes(paper),
                                  content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        data = SimpleUploadedFile("trial.csv", csv, content_type="text/csv")
        junk = SimpleUploadedFile("notes.zip", b"PK\x03\x04junk", content_type="application/zip")

        resp = self.client.post(self.URL, {"files": [docx, data, junk]}, format="multipart")
        self.assertEqual(resp.status_code, status.HTTP_200_OK, resp.content)
        body = resp.json()
        self.assertIn("verdict_distribution", body)
        self.assertIn("ingestion", body)
        self.assertEqual(body["ingestion"]["n_files"], 3)
        self.assertEqual(body["ingestion"]["n_data_files"], 1)
        self.assertEqual(body["ingestion"]["n_unknown"], 1)
        # the faithful claim should verify against its own data
        self.assertIn("VERIFIED", body["verdict_distribution"], body["verdict_distribution"])

    def test_empty_bundle_rejected(self):
        resp = self.client.post(self.URL, {}, format="multipart")
        self.assertEqual(resp.status_code, status.HTTP_400_BAD_REQUEST)

    def test_bundle_of_only_unknown_files_rejected(self):
        junk = SimpleUploadedFile("a.zip", b"PK\x03\x04junk", content_type="application/zip")
        resp = self.client.post(self.URL, {"files": [junk]}, format="multipart")
        self.assertEqual(resp.status_code, status.HTTP_400_BAD_REQUEST)
        self.assertIn("ingestion", resp.json())
